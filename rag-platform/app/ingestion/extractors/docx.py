"""DOCX extraction: text + structure (sections/headings)."""
from __future__ import annotations
from io import BytesIO
from typing import List, Tuple

from docx import Document
from docx.document import Document as DocxDocument
from docx.table import Table
from docx.text.paragraph import Paragraph


def extract_docx(data: bytes) -> Tuple[str, List[dict]]:
    """
    Extract full text and section structure (by headings).
    Returns (full_text, [{"section": "Heading 1", "text": "..."}, ...]).
    """
    doc = Document(BytesIO(data))
    sections = []
    full_parts = []
    current_section = None
    current_text = []
    for el in _iter_blocks(doc):
        if el["type"] == "heading":
            if current_section is not None and current_text:
                sections.append({"section": current_section, "text": "\n".join(current_text)})
                full_parts.append("\n".join(current_text))
            current_section = el["text"]
            current_text = []
        else:
            current_text.append(el["text"])
    if current_section is not None or current_text:
        sections.append({"section": current_section or "", "text": "\n".join(current_text)})
        full_parts.append("\n".join(current_text))
    if not sections:
        full_text = "\n".join(p.text for p in doc.paragraphs)
        return full_text, [{"section": "", "text": full_text}]
    full_text = "\n\n".join(full_parts)
    return full_text, sections


def _iter_blocks(doc: DocxDocument):
    for p in doc.paragraphs:
        style = (p.style and p.style.name or "").lower()
        if "heading" in style or p.text.strip().startswith("#"):
            yield {"type": "heading", "text": p.text.strip()}
        else:
            yield {"type": "paragraph", "text": p.text}
    for table in doc.tables:
        for row in table.rows:
            yield {"type": "paragraph", "text": " | ".join(c.text for c in row.cells)}
